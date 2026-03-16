#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
问卷分析工具 - 报告格式转换
============================

将 Markdown 报告转换为 TXT / Excel / Word 格式。

用法:
    # 转 Word
    python report_export.py --input report.md --format docx

    # 转 Excel
    python report_export.py --input report.md --format xlsx

    # 转 TXT
    python report_export.py --input report.md --format txt

    # 指定输出路径
    python report_export.py --input report.md --format docx --output my_report.docx

支持格式: md, txt, xlsx, docx
"""

import argparse
import json
import os
import re
import sys


# ========================================================================= #
#                        Markdown 解析器
# ========================================================================= #

def _parse_md_blocks(md_text: str) -> list:
    """
    将 Markdown 文本解析为结构化 block 列表。

    返回格式:
        [
            {"type": "h1", "text": "标题"},
            {"type": "h2", "text": "二级标题"},
            {"type": "h3", "text": "三级标题"},
            {"type": "blockquote", "text": "引用内容"},
            {"type": "table", "headers": [...], "rows": [...]},
            {"type": "paragraph", "text": "普通段落"},
            {"type": "list", "items": ["项1", "项2"]},
            {"type": "hr"},
        ]
    """
    blocks = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # 水平分割线
        if re.match(r'^---+$', stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 标题
        h_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            blocks.append({"type": f"h{level}", "text": text})
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "blockquote", "text": "\n".join(quote_lines)})
            continue

        # 表格
        if "|" in stripped and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].strip()):
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # 跳过表头和分隔线
            rows = []
            while i < len(lines) and "|" in lines[i]:
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # 列表
        if re.match(r'^[\-\*\d]+[\.\)]\s', stripped):
            items = []
            while i < len(lines) and re.match(r'^[\s]*[\-\*\d]+[\.\)]\s', lines[i]):
                item_text = re.sub(r'^[\s]*[\-\*\d]+[\.\)]\s+', '', lines[i])
                items.append(item_text.strip())
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        # 普通段落（合并连续非空行）
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(("#", ">", "|", "---")):
            para_lines.append(lines[i].strip())
            i += 1
            # 检查下一行是否是列表或表格的开始
            if i < len(lines):
                next_stripped = lines[i].strip()
                if re.match(r'^[\-\*\d]+[\.\)]\s', next_stripped):
                    break
                if "|" in next_stripped:
                    break
        blocks.append({"type": "paragraph", "text": "\n".join(para_lines)})

    return blocks


def _strip_md_inline(text: str) -> str:
    """去除 Markdown 行内格式（加粗、斜体、代码、链接、emoji等）"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # **bold**
    text = re.sub(r'\*(.+?)\*', r'\1', text)       # *italic*
    text = re.sub(r'`(.+?)`', r'\1', text)         # `code`
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text) # [text](url)
    return text


# ========================================================================= #
#                        TXT 导出
# ========================================================================= #

def export_txt(md_text: str, output_path: str) -> dict:
    """将 Markdown 转为纯文本"""
    blocks = _parse_md_blocks(md_text)
    lines = []

    for block in blocks:
        btype = block["type"]

        if btype == "hr":
            lines.append("=" * 60)
            lines.append("")

        elif btype.startswith("h"):
            level = int(btype[1])
            text = _strip_md_inline(block["text"])
            if level == 1:
                lines.append("=" * 60)
                lines.append(text)
                lines.append("=" * 60)
            elif level == 2:
                lines.append("")
                lines.append("-" * 40)
                lines.append(text)
                lines.append("-" * 40)
            else:
                lines.append("")
                lines.append(f"{'  ' * (level - 2)}【{text}】")
            lines.append("")

        elif btype == "blockquote":
            for line in block["text"].split("\n"):
                lines.append(f"    | {_strip_md_inline(line)}")
            lines.append("")

        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            # 计算每列最大宽度
            all_rows = [headers] + rows
            col_widths = []
            for ci in range(len(headers)):
                max_w = max(
                    len(_strip_md_inline(r[ci])) if ci < len(r) else 0
                    for r in all_rows
                )
                # 中文字符占两个宽度
                col_widths.append(max(max_w + 4, 8))

            def _fmt_row(row):
                cells = []
                for ci in range(len(headers)):
                    val = _strip_md_inline(row[ci]) if ci < len(row) else ""
                    cells.append(val.ljust(col_widths[ci]))
                return "  ".join(cells)

            lines.append(_fmt_row(headers))
            lines.append("  ".join("-" * w for w in col_widths))
            for row in rows:
                lines.append(_fmt_row(row))
            lines.append("")

        elif btype == "list":
            for item in block["items"]:
                lines.append(f"  • {_strip_md_inline(item)}")
            lines.append("")

        elif btype == "paragraph":
            lines.append(_strip_md_inline(block["text"]))
            lines.append("")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return {"status": "success", "output_path": output_path, "format": "txt"}


# ========================================================================= #
#                        Excel 导出
# ========================================================================= #

def export_xlsx(md_text: str, output_path: str) -> dict:
    """将 Markdown 报告转为结构化 Excel"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from _styles import (
        Theme, thin_border, make_fill,
        ALIGN_CENTER, ALIGN_LEFT, ALIGN_TOP_LEFT,
    )

    blocks = _parse_md_blocks(md_text)
    wb = Workbook()
    ws = wb.active
    ws.title = "分析报告"
    ws.sheet_view.showGridLines = False

    border = thin_border()
    row = 1

    for block in blocks:
        btype = block["type"]

        if btype == "hr":
            row += 1
            continue

        if btype.startswith("h"):
            level = int(btype[1])
            text = _strip_md_inline(block["text"])
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            cell = ws.cell(row=row, column=1, value=text)
            if level == 1:
                cell.font = Font(name="微软雅黑", size=16, bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="2B579A")
                cell.alignment = ALIGN_CENTER
                ws.row_dimensions[row].height = 42
            elif level == 2:
                cell.font = Font(name="微软雅黑", size=13, bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="4472C4")
                cell.alignment = ALIGN_LEFT
                ws.row_dimensions[row].height = 32
            elif level == 3:
                cell.font = Font(name="微软雅黑", size=11, bold=True, color="2B579A")
                cell.fill = PatternFill("solid", fgColor="D6E4F0")
                cell.alignment = ALIGN_LEFT
                ws.row_dimensions[row].height = 28
            else:
                cell.font = Font(name="微软雅黑", size=10, bold=True)
                cell.alignment = ALIGN_LEFT
                ws.row_dimensions[row].height = 24
            cell.border = border
            for c in range(2, 7):
                ws.cell(row=row, column=c).border = border
            row += 1

        elif btype == "blockquote":
            text = _strip_md_inline(block["text"])
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            cell = ws.cell(row=row, column=1, value=text)
            cell.font = Font(name="微软雅黑", size=9, italic=True, color="666666")
            cell.fill = PatternFill("solid", fgColor="F5F5F5")
            cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
            cell.border = border
            for c in range(2, 7):
                ws.cell(row=row, column=c).border = border
            line_count = max(1, len(text) // 80 + 1)
            ws.row_dimensions[row].height = max(28, line_count * 18)
            row += 1

        elif btype == "table":
            headers = block["headers"]
            rows_data = block["rows"]
            n_cols = len(headers)

            # 表头
            for ci, h in enumerate(headers, 1):
                cell = ws.cell(row=row, column=ci, value=_strip_md_inline(h))
                cell.font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="4472C4")
                cell.alignment = ALIGN_CENTER
                cell.border = border
            ws.row_dimensions[row].height = 28
            row += 1

            # 数据行
            for ri, data_row in enumerate(rows_data):
                fill = PatternFill("solid", fgColor="F2F7FB") if ri % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
                for ci in range(n_cols):
                    val = _strip_md_inline(data_row[ci]) if ci < len(data_row) else ""
                    cell = ws.cell(row=row, column=ci + 1, value=val)
                    cell.font = Font(name="微软雅黑", size=9)
                    cell.fill = fill
                    cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
                    cell.border = border
                ws.row_dimensions[row].height = 24
                row += 1
            row += 1  # 表格后空行

        elif btype == "list":
            for item in block["items"]:
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
                cell = ws.cell(row=row, column=1, value=f"  • {_strip_md_inline(item)}")
                cell.font = Font(name="微软雅黑", size=9)
                cell.alignment = ALIGN_TOP_LEFT
                ws.row_dimensions[row].height = 22
                row += 1
            row += 1

        elif btype == "paragraph":
            text = _strip_md_inline(block["text"])
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            cell = ws.cell(row=row, column=1, value=text)
            cell.font = Font(name="微软雅黑", size=9)
            cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
            line_count = max(1, len(text) // 80 + 1)
            ws.row_dimensions[row].height = max(22, line_count * 16)
            row += 1

    # 列宽
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 22
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 18
    ws.column_dimensions['E'].width = 18
    ws.column_dimensions['F'].width = 30

    wb.save(output_path)
    return {"status": "success", "output_path": output_path, "format": "xlsx"}


# ========================================================================= #
#                        Word 导出
# ========================================================================= #

def export_docx(md_text: str, output_path: str) -> dict:
    """将 Markdown 报告转为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        return {"error": "请安装 python-docx: pip install python-docx"}

    blocks = _parse_md_blocks(md_text)
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for block in blocks:
        btype = block["type"]

        if btype == "hr":
            # 添加一条分割线段落
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)

        elif btype.startswith("h"):
            level = int(btype[1])
            text = _strip_md_inline(block["text"])
            if level <= 3:
                heading = doc.add_heading(text, level=level)
                for run in heading.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        elif btype == "blockquote":
            text = _strip_md_inline(block["text"])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        elif btype == "table":
            headers = block["headers"]
            rows_data = block["rows"]
            n_cols = len(headers)
            n_rows = len(rows_data) + 1

            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Light Grid Accent 1"

            # 表头
            for ci, h in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = _strip_md_inline(h)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.name = "微软雅黑"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 数据行
            for ri, data_row in enumerate(rows_data):
                for ci in range(n_cols):
                    val = _strip_md_inline(data_row[ci]) if ci < len(data_row) else ""
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = val
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
                            run.font.name = "微软雅黑"
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 表后空行
            doc.add_paragraph()

        elif btype == "list":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(_strip_md_inline(item))
                run.font.size = Pt(10)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        elif btype == "paragraph":
            text = _strip_md_inline(block["text"])
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.save(output_path)
    return {"status": "success", "output_path": output_path, "format": "docx"}


# ========================================================================= #
#                        主入口
# ========================================================================= #

FORMAT_MAP = {
    "txt": export_txt,
    "xlsx": export_xlsx,
    "docx": export_docx,
    "word": export_docx,
    "excel": export_xlsx,
}


def convert_report(input_path: str, fmt: str, output_path: str = None) -> dict:
    """
    将 Markdown 报告转换为指定格式。

    Args:
        input_path: 输入的 .md 文件路径
        fmt: 目标格式 (txt / xlsx / docx / word / excel)
        output_path: 输出路径（可选，默认替换扩展名）

    Returns:
        {"status": "success", "output_path": str, "format": str}
    """
    fmt = fmt.lower().strip(".")

    if fmt == "md" or fmt == "markdown":
        return {"status": "success", "output_path": input_path, "format": "md", "message": "已经是 Markdown 格式"}

    if fmt not in FORMAT_MAP:
        return {"error": f"不支持的格式: {fmt}。支持: md, txt, xlsx, docx"}

    if not os.path.exists(input_path):
        return {"error": f"文件不存在: {input_path}"}

    with open(input_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    if not output_path:
        ext = "docx" if fmt == "word" else ("xlsx" if fmt == "excel" else fmt)
        base = os.path.splitext(input_path)[0]
        output_path = f"{base}.{ext}"

    return FORMAT_MAP[fmt](md_text, output_path)


def main():
    parser = argparse.ArgumentParser(description="问卷分析报告格式转换")
    parser.add_argument("--input", required=True, help="输入的 Markdown 报告路径")
    parser.add_argument("--format", required=True, help="目标格式: txt / xlsx / docx / word / excel")
    parser.add_argument("--output", default=None, help="输出文件路径（可选，默认自动生成）")
    args = parser.parse_args()

    result = convert_report(args.input, args.format, args.output)
    print(json.dumps(result, ensure_ascii=False, indent=2))

    if "error" in result:
        sys.exit(1)


if __name__ == "__main__":
    main()
